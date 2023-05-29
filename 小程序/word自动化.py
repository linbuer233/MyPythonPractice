import os
import sys

from docx import Document
from docx.oxml.ns import qn
from docx.oxml.shared import OxmlElement


def insertHR(paragraph, color, width):
    p = paragraph._p  # p is the <w:p> XML element
    pPr = p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    pPr.insert_element_before(pBdr,
                              'w:shd', 'w:tabs', 'w:suppressAutoHyphens', 'w:kinsoku', 'w:wordWrap',
                              'w:overflowPunct', 'w:topLinePunct', 'w:autoSpaceDE', 'w:autoSpaceDN',
                              'w:bidi', 'w:adjustRightInd', 'w:snapToGrid', 'w:spacing', 'w:ind',
                              'w:contextualSpacing', 'w:mirrorIndents', 'w:suppressOverlap', 'w:jc',
                              'w:textDirection', 'w:textAlignment', 'w:textboxTightWrap',
                              'w:outlineLvl', 'w:divId', 'w:cnfStyle', 'w:rPr', 'w:sectPr',
                              'w:pPrChange'
                              )
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'thick')
    bottom.set(qn('w:sz'), width)
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), color)
    pBdr.append(bottom)


filepath = sys.argv[1]
for _ in range(1):
    # doc 转 docx，方便 python-docx 处理
    from win32com.client import Dispatch

    word = Dispatch('Word.Application')
    # doc 转化 docx
    doc = word.Documents.Open(filepath)  # ('D:\\micaps\\text\\贺州市 - 旬月天气预报 -2022011-2022 年 4 月下旬天气预报.doc')
    docxNamePath = os.path.dirname(filepath) + '\\text.docx'
    doc.SaveAs(docxNamePath, 12)  # , False, "", True, "", False, False, False, False)
    doc.Close()
    word.Quit()

    document = Document(os.path.dirname(filepath) + '\\text.docx')
    # print(dir(document.styles['Normal']._element.rPr.rFonts))#.set(qn("w:eastAsia"),"宋体")
    count = 1
    bra = 0
    temp = 0
    temp1 = 0
    a = []
    num = -1
    for i, paragraph in enumerate(document.paragraphs):
        for j, run in enumerate(paragraph.runs):
            if run.text == '贺州市气象局' and len(paragraph.runs) > 2:
                count = 0
                continue
            if count == 0:
                if run.text == '签发人：':
                    bra = 1
                    break
                a.append(len(run.text))
        if bra == 1:
            break
    for i, paragraph in enumerate(document.paragraphs):
        for j, run in enumerate(paragraph.runs):
            # run._element.rPr.rFonts.set(qn("w:eastAsia"),"宋体")
            if count == 0 and run.text == '贺州市气象局':
                run_text = run.text.replace('贺州市气象局', '富川瑶族自治县气象局')
                run.text = run_text
                insertHR(paragraph, 'red', '18')
                count = 1
                continue
            if count == 1:
                num += 1
                run_text = run.text.replace(run.text, ' ' * int(a[num] / 40 * 29))
                run.text = run_text
                if num == len(a) - 1:
                    count = 2
            run_text = run.text.replace('贺州市', '富川县')
            run.text = run_text
            run_text = run.text.replace('州市', '川县')
            run.text = run_text
            run_text = run.text.replace('贺', '富')
            run.text = run_text
            run_text = run.text.replace('贺州市气象局', '富川县气象局')
            run.text = run_text
            if (run.text == '黎') or (temp == 1):
                temp1 += 1
                temp = 1
                run_text = run.text.replace('黎', '李')
                run.text = run_text
                run_text = run.text.replace('  ', '付')
                run.text = run_text
                run_text = run.text.replace('萍', '宇')
                run.text = run_text
                if temp1 == 3:
                    temp = 0
                continue
            run_text = run.text.replace('黎  萍', '李付宇')
            run.text = run_text
            run_text = run.text.replace('市', '县')
            run.text = run_text
            if i == len(document.paragraphs) - 2:
                run_text = run.text.replace('5222096', '7882338')
                run.text = run_text
                insertHR(paragraph, 'black', '10')
    outname = os.path.split(filepath)[-1]
    outname = outname.replace('贺州市', '富川县')
    outname = outname.replace('doc', 'docx')
    document.save(os.path.dirname(filepath) + '\\' + outname)

os.system('rm ' + os.path.dirname(filepath) + '\\text.docx')
